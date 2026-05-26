"""
Contract review service.

Pipeline:
  1. extract_text()  — PDF/DOCX → plain text (server-side, no external storage)
  2. analyze_contract() — plain text → structured JSON via Qwen LLM
                        — optional FaRui supplementary opinion (non-fatal)
"""
import io
import json
import logging
import re
from typing import Any, Optional

import pdfplumber
from docx import Document as DocxDocument
from openai import OpenAI


logger = logging.getLogger(__name__)

# Truncate to ~7 000 tokens of English text — fits comfortably inside the
# qwen-plus 32k context window along with the system prompt.
MAX_CHARS = 28_000

# ---------------------------------------------------------------------------
# System prompt
# ---------------------------------------------------------------------------

CONTRACT_ANALYSIS_SYSTEM_PROMPT = """\
You are a senior commercial law attorney with dual expertise in Chinese law \
(PRC Civil Code Book III Contract Section — 民法典合同编, foreign investment law, company law) \
and international contract practice (UAE law, CISG, ICC rules).

Your task: review the contract text supplied by the user and return a \
structured JSON risk assessment.

Instructions:
1. Identify ALL parties by name, their contractual role (Party A / Party B / \
   Guarantor / etc.), and their registered jurisdiction.
2. Determine the governing law and dispute resolution jurisdiction.
3. Assign an overall risk score 0–100 (higher = more risky to the weaker \
   party) and a risk level: LOW (0–29), MEDIUM (30–59), HIGH (60–79), \
   CRITICAL (80–100).
4. Analyze every significant clause individually. For each clause state:
   - clause name
   - risk level (LOW / MEDIUM / HIGH / CRITICAL)
   - the specific legal issue or imbalance
   - a concrete revision suggestion (plain language, not legalese)
5. List clauses that are MISSING but should be present under standard \
   commercial practice and/or Chinese law requirements.
6. List clauses that pose CRITICAL risk and require immediate attention.
7. Provide top-level actionable suggestions ranked by severity \
   (CRITICAL first, then HIGH, MEDIUM, LOW).
8. Write a concise executive summary in plain English (3–5 sentences).

Apply Chinese law perspective throughout:
- Flag clauses that may be unenforceable under PRC Civil Code (e.g., \
  unlimited liability waivers, grossly one-sided penalty clauses, \
  IP ownership defaults that contradict Chinese mandatory rules).
- Note foreign exchange control implications (SAFE regulations) if a \
  Chinese entity is or may become a party.
- Identify cross-border enforceability gaps between UAE and Chinese courts.
- Flag missing GDPR / PIPL (Personal Information Protection Law — 个人信息保护法) \
  data-privacy provisions if personal data is involved. Always refer to laws in English.

CRITICAL INSTRUCTION: Return ONLY a single valid JSON object — no markdown \
code fences, no preamble text, no trailing commentary.  The JSON must \
conform exactly to this schema:

{
  "riskScore": <integer 0–100>,
  "riskLevel": "<LOW|MEDIUM|HIGH|CRITICAL>",
  "summary": "<executive summary>",
  "parties": [
    {"name": "<string>", "role": "<string>", "jurisdiction": "<string>"}
  ],
  "governingLaw": "<string>",
  "jurisdiction": "<string>",
  "clauseAnalysis": [
    {
      "clause": "<clause name>",
      "risk": "<LOW|MEDIUM|HIGH|CRITICAL>",
      "issue": "<description of legal problem>",
      "suggestion": "<concrete revision recommendation>"
    }
  ],
  "missingClauses": ["<clause name>", ...],
  "criticalClauses": ["<clause name>", ...],
  "suggestions": [
    {"type": "<LOW|MEDIUM|HIGH|CRITICAL>", "text": "<actionable suggestion>"}
  ]
}
"""


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Returns cleaned plain text truncated to MAX_CHARS.
    Raises ValueError for unsupported formats or unreadable files.
    """
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        raw = _extract_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        raw = _extract_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '{filename}'. Only .pdf and .docx files are accepted."
        )

    text = _clean_text(raw)
    if not text.strip():
        raise ValueError(
            "No readable text could be extracted from the uploaded file. "
            "Make sure it is not a scanned image PDF."
        )
    return text[:MAX_CHARS]


def _extract_pdf(file_bytes: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if page_text:
                parts.append(page_text)
    return "\n\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def _clean_text(text: str) -> str:
    # Collapse 3+ consecutive blank lines → 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip non-printable characters except tab, LF, CR, and standard printable range
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", "", text)
    return text.strip()


# ---------------------------------------------------------------------------
# LLM analysis
# ---------------------------------------------------------------------------

def analyze_contract(
    contract_text: str,
    config: dict[str, Any],
    farui_service: Optional[Any] = None,
) -> dict[str, Any]:
    """
    Run contract review via Qwen LLM.  Optionally attaches a FaRui legal
    opinion as a non-fatal supplementary field.

    Returns a dict conforming to the CONTRACT_ANALYSIS_SYSTEM_PROMPT schema,
    always with safe defaults for every field.

    Raises RuntimeError on unrecoverable LLM failure.
    """
    api_key = config.get("LLM_API_KEY", "")
    base_url = config.get("LLM_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")
    model = (
        config.get("CONTRACT_ANALYSIS_MODEL")
        or config.get("LLM_CHAT_MODEL")
        or "qwen-plus"
    )

    client = OpenAI(api_key=api_key, base_url=base_url)

    user_message = (
        "Please analyze the following contract and return a structured JSON risk assessment.\n\n"
        "CONTRACT TEXT:\n"
        "---\n"
        f"{contract_text}\n"
        "---"
    )

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": CONTRACT_ANALYSIS_SYSTEM_PROMPT},
                {"role": "user", "content": user_message},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        raw_json = response.choices[0].message.content or "{}"
    except Exception as exc:
        logger.error("Contract analysis LLM call failed: %s", exc)
        raise RuntimeError(f"Contract analysis failed: {type(exc).__name__}: {exc}") from exc

    analysis = _parse_and_validate(raw_json)

    # --- Optional FaRui supplementary legal opinion ---
    if farui_service is not None:
        try:
            snippet = contract_text[:3_000]
            farui_result = farui_service.reply(
                message=(
                    "Please provide a brief legal opinion (5–7 bullet points in English) "
                    "on the key legal risks and recommendations for this contract:\n\n"
                    f"{snippet}"
                ),
                history=[],
                deep_think=False,
                online_search=False,
            )
            analysis["faruiOpinion"] = farui_result.get("reply") or None
        except Exception as exc:  # noqa: BLE001
            logger.warning("FaRui supplementary opinion failed (non-fatal): %s", exc)
            analysis["faruiOpinion"] = None
    else:
        analysis["faruiOpinion"] = None

    return analysis


def _parse_and_validate(raw: str) -> dict[str, Any]:
    """
    Parse LLM JSON output and apply safe defaults for every schema field.
    This ensures the frontend always receives a consistent payload even when
    the LLM omits optional fields.
    """
    # Strip accidental markdown fences the model might add despite instructions
    cleaned = re.sub(r"^```(?:json)?\s*", "", raw.strip(), flags=re.MULTILINE)
    cleaned = re.sub(r"\s*```$", "", cleaned.strip(), flags=re.MULTILINE)

    try:
        data: dict[str, Any] = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error("LLM returned non-JSON: %s", raw[:500])
        raise RuntimeError("LLM returned a malformed JSON response.") from exc

    # Clamp riskScore
    score = data.get("riskScore", 50)
    if isinstance(score, (int, float)):
        score = max(0, min(100, int(score)))
    else:
        score = 50
    data["riskScore"] = score

    # Apply defaults
    data.setdefault("riskLevel", _derive_risk_level(score))
    data.setdefault("summary", "Analysis summary was not generated.")
    data.setdefault("parties", [])
    data.setdefault("governingLaw", "Not specified")
    data.setdefault("jurisdiction", "Not specified")
    data.setdefault("clauseAnalysis", [])
    data.setdefault("missingClauses", [])
    data.setdefault("criticalClauses", [])
    data.setdefault("suggestions", [])

    return data


def _derive_risk_level(score: int) -> str:
    if score >= 80:
        return "CRITICAL"
    if score >= 60:
        return "HIGH"
    if score >= 30:
        return "MEDIUM"
    return "LOW"
